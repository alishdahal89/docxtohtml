from docx import Document

outputDoc = Document()
Html_file= open("outputHTML","w")
Html_image_file= open("outputImageTags","w")
inputDoc = Document('Article.docx')

# Image tagging variables
curriculum = "CHEM"
url = "https://j2ids9z9x.cloudfront.net/diagrams/"
unit = 1
section = 0
imageNumber = 0

for paragraph in inputDoc.paragraphs:
    if len(paragraph.text) == 0:
        continue
    paragraphText = []
    paragraphText.append("<p>")

    # Check if run contains combination of bold, italic, and underline and append tags accordingly if detected
    for run in paragraph.runs:
        boldTag = False
        italicTag = False
        underlineTag = False
        if run.bold:
            boldTag = True
            paragraphText.append("<b>")
        if run.italic:
            italicTag = True
            paragraphText.append("<i>")
        if run.underline:
            underlineTag = True
            paragraphText.append("<u>")
        paragraphText.append(run.text)
        if boldTag:
            paragraphText.append("</b>")
        if italicTag:
            paragraphText.append("</i>")
        if underlineTag:
            paragraphText.append("</u>")

    paragraphText.append("</p>")
    outputDoc.add_paragraph("".join(paragraphText))
    Html_file.write("".join(paragraphText) + "\n")

for inline_shape in inputDoc.inline_shapes: # Image checking
    Html_image_file.write("<img src = \"" + url + curriculum + "-" + str(unit) + "-" + str(section) + "-X-" + str(imageNumber) + ".png\"><br>" + "\n")
    imageNumber += 1

outputDoc.save('outputDOCX.docx')
Html_file.close()
